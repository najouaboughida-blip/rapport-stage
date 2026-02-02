"""
Générateur de documents Word académiques
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os
import tempfile
from typing import Dict, List, Optional
from io import BytesIO
import re

class WordGenerator:
    """Générateur de rapports Word académiques"""
    
    def __init__(self, config: Optional[Dict] = None):
        """
        Initialise le générateur Word
        
        Args:
            config: Configuration Word
        """
        self.config = config or {
            'font_name': 'Times New Roman',
            'title_font_size': 16,
            'heading1_font_size': 14,
            'heading2_font_size': 12,
            'body_font_size': 11,
            'line_spacing': 1.5,
            'margins': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5}  # cm
        }
        
        self.document = Document()
        self._setup_document()
        self._setup_styles()
    
    def _setup_document(self):
        """Configure le document Word"""
        # Marges
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(self.config['margins']['top'])
            section.bottom_margin = Cm(self.config['margins']['bottom'])
            section.left_margin = Cm(self.config['margins']['left'])
            section.right_margin = Cm(self.config['margins']['right'])
            
            # En-têtes et pieds de page
            header = section.header
            footer = section.footer
            
            # En-tête minimal
            header_para = header.paragraphs[0]
            header_para.text = ""
            
            # Pied de page avec numéro de page
            footer_para = footer.paragraphs[0]
            footer_para.text = ""
    
    def _setup_styles(self):
        """Configure les styles académiques"""
        styles = self.document.styles
        
        # Style Titre principal
        if 'AcademicTitle' not in styles:
            title_style = styles.add_style('AcademicTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.config['font_name']
            title_style.font.size = Pt(self.config['title_font_size'])
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)  # #2C3E50
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)
            title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Style Titre 1
        if 'AcademicHeading1' not in styles:
            heading1_style = styles.add_style('AcademicHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = self.config['font_name']
            heading1_style.font.size = Pt(self.config['heading1_font_size'])
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(44, 62, 80)
            heading1_style.paragraph_format.space_before = Pt(20)
            heading1_style.paragraph_format.space_after = Pt(10)
            heading1_style.paragraph_format.keep_with_next = True
        
        # Style Titre 2
        if 'AcademicHeading2' not in styles:
            heading2_style = styles.add_style('AcademicHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = self.config['font_name']
            heading2_style.font.size = Pt(self.config['heading2_font_size'])
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(52, 73, 94)  # #34495E
            heading2_style.paragraph_format.space_before = Pt(15)
            heading2_style.paragraph_format.space_after = Pt(8)
            heading2_style.paragraph_format.keep_with_next = True
        
        # Style Texte normal académique
        if 'AcademicNormal' not in styles:
            normal_style = styles.add_style('AcademicNormal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = self.config['font_name']
            normal_style.font.size = Pt(self.config['body_font_size'])
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            normal_style.paragraph_format.line_spacing = self.config['line_spacing']
            normal_style.paragraph_format.first_line_indent = Cm(0.63)  # 0.63cm = 0.25 inch
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Style pour les listes
        if 'AcademicList' not in styles:
            list_style = styles.add_style('AcademicList', WD_STYLE_TYPE.PARAGRAPH)
            list_style.font.name = self.config['font_name']
            list_style.font.size = Pt(self.config['body_font_size'])
            list_style.paragraph_format.left_indent = Cm(1.25)
            list_style.paragraph_format.first_line_indent = Cm(-0.63)
            list_style.paragraph_format.space_after = Pt(3)
        
        # Style Footer
        if 'AcademicFooter' not in styles:
            footer_style = styles.add_style('AcademicFooter', WD_STYLE_TYPE.PARAGRAPH)
            footer_style.font.name = self.config['font_name']
            footer_style.font.size = Pt(9)
            footer_style.font.color.rgb = RGBColor(149, 165, 166)  # #95A5A6
            footer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, report_data: Dict, output_path: Optional[str] = None) -> str:
        """
        Génère un document Word
        
        Args:
            report_data: Données du rapport
            output_path: Chemin de sortie
        
        Returns:
            Chemin du fichier généré
        """
        # Créer un nom de fichier
        if not output_path:
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            student_name = report_data.get('student', {}).get('full_name', 'rapport').replace(' ', '_')
            filename = f"rapport_{student_name}_{timestamp}.docx"
            output_path = os.path.join(temp_dir, filename)
        
        # Réinitialiser le document
        self.document = Document()
        self._setup_document()
        self._setup_styles()
        
        # Ajouter les sections
        self._add_cover_page(report_data)
        self.document.add_page_break()
        
        # Remerciements
        if 'thanks' in report_data.get('sections', {}):
            self._add_section('REMERCIEMENTS', report_data['sections']['thanks'])
            self.document.add_page_break()
        
        # Résumés
        if 'abstract' in report_data.get('sections', {}):
            self._add_section('RÉSUMÉS', report_data['sections']['abstract'])
            self.document.add_page_break()
        
        # Table des matières (à générer manuellement pour l'instant)
        self._add_table_of_contents(report_data)
        self.document.add_page_break()
        
        # Sections principales
        sections_order = [
            ('introduction', 'INTRODUCTION GÉNÉRALE'),
            ('company_presentation', 'PRÉSENTATION DE L\'ENTREPRISE'),
            ('methodology', 'MÉTHODOLOGIE DE TRAVAIL'),
            ('realization', 'RÉALISATION TECHNIQUE'),
            ('results', 'RÉSULTATS ET DISCUSSION'),
            ('conclusion', 'CONCLUSION ET PERSPECTIVES'),
            ('bibliography', 'BIBLIOGRAPHIE'),
            ('annexes', 'ANNEXES')
        ]
        
        for section_key, section_title in sections_order:
            if section_key in report_data.get('sections', {}):
                self._add_section(section_title, report_data['sections'][section_key])
                if section_key != 'annexes':
                    self.document.add_page_break()
        
        # Sauvegarder
        self.document.save(output_path)
        print(f"✅ Word généré: {output_path}")
        
        return output_path
    
    def _add_cover_page(self, report_data: Dict):
        """Ajoute la page de couverture"""
        student = report_data.get('student', {})
        company = report_data.get('company', {})
        
        # Espacement initial
        for _ in range(6):
            self.document.add_paragraph()
        
        # Université
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.document.styles['AcademicTitle']
        run = p.add_run("UNIVERSITÉ MOHAMMED PREMIER")
        run.font.size = Pt(14)
        
        # École
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("École Nationale des Sciences Appliquées - Oujda")
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Filière
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student.get('filiere', 'Génie Informatique'))
        run.font.name = self.config['font_name']
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Espacement
        for _ in range(4):
            self.document.add_paragraph()
        
        # Titre principal
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.document.styles['AcademicTitle']
        run = p.add_run("RAPPORT DE STAGE DE FIN D'ÉTUDES")
        run.font.size = Pt(16)
        
        # Sous-titre
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Présenté en vue de l'obtention du Diplôme d'Ingénieur d'État")
        run.font.name = self.config['font_name']
        run.font.size = Pt(11)
        
        # Espacement
        for _ in range(4):
            self.document.add_paragraph()
        
        # Titre du projet
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{student.get("project_title", "Titre du projet")}"')
        run.font.name = self.config['font_name']
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(41, 128, 185)  # #2980B9
        run.italic = True
        
        # Espacement
        for _ in range(5):
            self.document.add_paragraph()
        
        # Informations dans un tableau centré
        table = self.document.add_table(rows=5, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Configurer la largeur
        for row in table.rows:
            row.cells[0].width = Cm(12)
        
        # Remplir les informations
        rows = table.rows
        
        # Étudiant
        cell = rows[0].cells[0]
        cell.text = "Présenté par :"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student.get('full_name', 'NOM Prénom'))
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Encadrants
        cell = rows[1].cells[0]
        cell.text = "Encadré par :"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{student.get('supervisor', 'Dr. NOM Prénom')} (ENSAO)")
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{company.get('supervisor', 'M. NOM Prénom')} ({company.get('name', 'Entreprise')})")
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Entreprise
        cell = rows[2].cells[0]
        cell.text = "Entreprise d'accueil :"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company.get('name', 'Nom Entreprise'))
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Période
        cell = rows[3].cells[0]
        cell.text = "Période de stage :"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student.get('duration', '2 mois'))
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Année
        cell = rows[4].cells[0]
        cell.text = "Année universitaire :"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student.get('academic_year', '2024-2025'))
        run.font.name = self.config['font_name']
        run.font.size = Pt(12)
        
        # Espacement
        for _ in range(6):
            self.document.add_paragraph()
        
        # Footer
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = self.document.styles['AcademicFooter']
        p.add_run("Mémoire de Projet de Fin d'Études - Génie Informatique - ENSAO")
    
    def _add_section(self, title: str, content: str):
        """Ajoute une section au document"""
        # Titre
        p = self.document.add_paragraph(title, style='AcademicHeading1')
        
        # Contenu
        self._add_html_content(content)
    
    def _add_html_content(self, html: str):
        """Ajoute du contenu HTML formaté"""
        if not html:
            return
        
        # Nettoyer le HTML
        html = self._clean_html(html)
        
        # Traiter les paragraphes
        paragraphs = re.split(r'</?p[^>]*>', html)
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Détecter les titres
            if para_text.startswith('<h2>'):
                title = para_text.replace('<h2>', '').replace('</h2>', '').strip()
                self.document.add_paragraph(title, style='AcademicHeading2')
            elif para_text.startswith('<h3>'):
                title = para_text.replace('<h3>', '').replace('</h3>', '').strip()
                self.document.add_paragraph(title, style='Heading 3')
            elif para_text.startswith('<ul>') or '<li>' in para_text:
                # Gérer les listes
                list_items = re.findall(r'<li>(.*?)</li>', para_text, re.DOTALL)
                for item in list_items:
                    item = item.strip()
                    if item:
                        p = self.document.add_paragraph(style='AcademicList')
                        p.add_run(f"• {item}")
            else:
                # Texte normal
                para_text = re.sub(r'<[^>]+>', '', para_text)
                para_text = para_text.replace('&nbsp;', ' ')
                para_text = para_text.replace('&amp;', '&')
                para_text = para_text.replace('&lt;', '<')
                para_text = para_text.replace('&gt;', '>')
                
                if para_text.strip():
                    p = self.document.add_paragraph(style='AcademicNormal')
                    p.add_run(para_text.strip())
    
    def _clean_html(self, html: str) -> str:
        """Nettoie le HTML pour Word"""
        if not html:
            return ""
        
        # Remplacer les balises de saut de ligne
        html = html.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        
        # Normaliser les espaces
        html = re.sub(r'\s+', ' ', html)
        
        return html.strip()
    
    def _add_table_of_contents(self, report_data: Dict):
        """Ajoute une table des matières"""
        p = self.document.add_paragraph("TABLE DES MATIÈRES", style='AcademicHeading1')
        
        # Sections à inclure
        sections = [
            ('REMERCIEMENTS', 1),
            ('RÉSUMÉS', 2),
            ('INTRODUCTION GÉNÉRALE', 3),
            ('PRÉSENTATION DE L\'ENTREPRISE', 4),
            ('MÉTHODOLOGIE DE TRAVAIL', 5),
            ('RÉALISATION TECHNIQUE', 6),
            ('RÉSULTATS ET DISCUSSION', 7),
            ('CONCLUSION ET PERSPECTIVES', 8),
            ('BIBLIOGRAPHIE', 9),
            ('ANNEXES', 10)
        ]
        
        for section_title, page_num in sections:
            # Vérifier si la section existe
            section_key = section_title.lower().replace(' ', '_').replace('é', 'e').replace("'", '')
            if section_key in report_data.get('sections', {}):
                p = self.document.add_paragraph()
                
                # Titre avec points de suite
                run_title = p.add_run(section_title)
                run_title.font.name = self.config['font_name']
                run_title.font.size = Pt(11)
                
                # Ajouter des points
                tab_positions = p.paragraph_format.tab_stops
                if not tab_positions:
                    p.paragraph_format.tab_stops.add_tab_stop(Cm(14))
                
                p.add_run('\t')
                
                # Numéro de page
                run_page = p.add_run(str(page_num))
                run_page.font.name = self.config['font_name']
                run_page.font.size = Pt(11)
    
    def generate_from_html(self, html_content: str, output_path: str, metadata: Dict = None) -> str:
        """
        Génère un document Word directement à partir de HTML
        
        Args:
            html_content: Contenu HTML
            output_path: Chemin de sortie
            metadata: Métadonnées
        
        Returns:
            Chemin du fichier généré
        """
        # Réinitialiser
        self.document = Document()
        self._setup_document()
        self._setup_styles()
        
        # Ajouter le contenu
        self._add_html_content(html_content)
        
        # Sauvegarder
        self.document.save(output_path)
        
        return output_path
    
    def generate_bytes(self, report_data: Dict) -> bytes:
        """
        Génère un document Word en mémoire
        
        Args:
            report_data: Données du rapport
        
        Returns:
            Document Word sous forme de bytes
        """
        # Générer dans un buffer
        buffer = BytesIO()
        
        # Réinitialiser
        self.document = Document()
        self._setup_document()
        self._setup_styles()
        
        # Ajouter la page de couverture
        self._add_cover_page(report_data)
        
        # Ajouter les sections principales
        sections_order = [
            ('introduction', 'INTRODUCTION GÉNÉRALE'),
            ('company_presentation', 'PRÉSENTATION DE L\'ENTREPRISE'),
            ('conclusion', 'CONCLUSION')
        ]
        
        for section_key, section_title in sections_order:
            if section_key in report_data.get('sections', {}):
                self._add_section(section_title, report_data['sections'][section_key])
                self.document.add_page_break()
        
        # Sauvegarder dans le buffer
        self.document.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()


def generate_quick_docx(content: str, title: str = "Rapport", student_name: str = "") -> bytes:
    """
    Génère un document Word simple rapidement
    
    Args:
        content: Contenu du document
        title: Titre du document
        student_name: Nom de l'auteur
    
    Returns:
        Document Word sous forme de bytes
    """
    doc = Document()
    
    # Titre
    doc.add_heading(title, 0)
    
    # Métadonnées
    if student_name:
        p = doc.add_paragraph()
        p.add_run(f"Auteur: {student_name}\n")
        p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}\n")
    
    # Séparateur
    doc.add_paragraph().add_run("=" * 50)
    
    # Contenu
    doc.add_paragraph(content)
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()